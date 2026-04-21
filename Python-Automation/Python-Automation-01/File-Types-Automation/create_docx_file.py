from docx import Document
from docx.shared import Pt

document = Document()
document.add_heading("JLoka-01 Heading")
p = document.add_paragraph("JLoka-02 Paragraph in bold and italic")
# p.add_run('bold').bold = True
# p.add_run('italic').italic = True
p.runs[0].bold = True
p.runs[0].italic = True
document.add_heading("JLoka-03 Test Automate docx files", level=2)
document.add_paragraph("Build JLoka-04 Automation Tools", style='Heading 3')
document.add_paragraph("JLoka-05 is ITE Developer", style='Quote')
############## Adding List Number ##############
document.add_paragraph("Automate PDF Files", style='List Number')
document.add_paragraph("Compare between different packages", style='List Number')
############## Adding List Bullet ##############
document.add_paragraph("Automate docx files", style='List Bullet')
document.add_paragraph("Check python-docx library", style='List Bullet')
############## Adding Picture ##############
document.add_picture("tomatoes.jpg", width=Pt(540), height=Pt(540))
############## Save Document ##############
document.save('jloka-test-docx.docx')